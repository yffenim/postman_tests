import os
import json
import logging
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(filename='postman_requests.log', level=logging.INFO, format='%(asctime)s %(levelname)s: %(message)s')

def extract_access_token(headers):
    for header in headers:
        if header['key'].lower() == 'accesstoken':
            return header['value']
    return None

def process_requests(collection_items, results):
    for item in collection_items:
        if 'request' in item:
            print(f"Requesting... {item['name']}")
            request = item['request']
            url = request['url']['raw']
            method = request['method']
            headers = request.get('header', [])
            access_token = extract_access_token(headers)

            custom_headers = {}
            if access_token:
                custom_headers['accesstoken'] = access_token

            try:
                response = requests.request(method, url, headers=custom_headers)
                results[f"{method} {item['name']}"] = (response.status_code, response.reason)
                logging.info(f"{method} {item['name']}: {response.status_code} - {response.reason}")
            except requests.exceptions.RequestException as e:
                error_message = str(e)
                if "Invalid URL" in error_message:
                    results[f"{method} {item['name']}"] = ('Error - Invalid URL', error_message)
                else:
                    results[f"{method} {item['name']}"] = ('Error', error_message)
                logging.error(f"Request failed for '{item['name']}' - {error_message}")
        if 'item' in item:
            process_requests(item['item'], results)


def create_word_document(results, title, doc):
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4)

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Request'
    header_cells[1].text = 'Status'
    header_cells[2].text = 'Active?'
    header_cells[3].text = 'Notes'

    for key, value in results.items():
        row_cells = table.add_row().cells
        is_not_found_or_invalid = value[0] == 404 or value[0] == 'Error - Invalid URL'

        row_cells[0].text = key
        row_cells[1].text = f"{value[0]} - {value[1]}"
        row_cells[2].text = ''
        row_cells[3].text = ''

        if is_not_found_or_invalid:
            for cell in row_cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()


def main():
    folder_path = './collections_json/'
    json_file_paths = [os.path.join(folder_path, file) for file in os.listdir(folder_path) if file.endswith('.json')]
    
    # Create and set word page to landscape 
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    for json_file_path in json_file_paths:
        with open(json_file_path, 'r') as file:
            postman_collection = json.load(file)

        title = postman_collection.get('info', {}).get('name', 'Untitled Collection')

        results = {}
        process_requests(postman_collection['item'], results)

        create_word_document(results, title, doc)

        for key, value in results.items():
            print(f"{key}: {value[0]} - {value[1]}")

    doc.save('postman_results.docx')

if __name__ == '__main__':
    main()
